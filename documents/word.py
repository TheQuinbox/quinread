from .base import BaseDocument
import docx

class DocxDocument(BaseDocument):
	def __init__(self, path):
		self.path = path
	
	def read(self):
		self.document = docx.Document(self.path)
		full_text = [para.text for para in self.document.paragraphs]
		return "\n".join(full_text)
	
	def close(self):
		super().close()
